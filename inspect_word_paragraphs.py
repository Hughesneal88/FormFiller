import docx

doc = docx.Document(r"C:\Users\Hughe\Downloads\survey_auto_filler\Data Analysis.docx")

with open("paragraphs.txt", "w", encoding="utf-8") as f:
    table_counter = 0
    for idx, element in enumerate(doc.element.body):
        if element.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(element, doc)
            if p.text.strip():
                f.write(f"[P {idx}]: {p.text}\n")
        elif element.tag.endswith('tbl'):
            f.write(f"[TBL {table_counter}]: (Table with {len(doc.tables[table_counter].rows)} rows and {len(doc.tables[table_counter].columns)} columns)\n")
            table_counter += 1
