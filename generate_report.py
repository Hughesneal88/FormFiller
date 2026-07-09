import os
import shutil
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Paths
EXCEL_PATH = r"C:\Users\Hughe\Downloads\survey_auto_filler\Artisanal_Fisher_Survey_Jamestown_Fishing_Harbour_all_versions_English.xlsx"
DOCX_PATH = r"C:\Users\Hughe\Downloads\survey_auto_filler\Data Analysis.docx"
DOCX_OUT_PATH = r"C:\Users\Hughe\Downloads\survey_auto_filler\Data Analysis_Fisheries.docx"
CHARTS_DIR = r"C:\Users\Hughe\Downloads\survey_auto_filler\temp_charts"

# Create charts directory if it doesn't exist
os.makedirs(CHARTS_DIR, exist_ok=True)

# 2. Premium Matplotlib Styling
def apply_premium_style():
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['font.sans-serif'] = ['Arial', 'Liberation Sans', 'DejaVu Sans']
    plt.rcParams['figure.facecolor'] = 'white'
    plt.rcParams['axes.facecolor'] = 'white'
    plt.rcParams['axes.edgecolor'] = '#cccccc'
    plt.rcParams['axes.grid'] = True
    plt.rcParams['grid.color'] = '#f0f0f0'
    plt.rcParams['grid.linestyle'] = '-'
    plt.rcParams['grid.alpha'] = 0.8
    plt.rcParams['text.color'] = '#2c3e50'
    plt.rcParams['axes.labelcolor'] = '#2c3e50'
    plt.rcParams['xtick.color'] = '#555555'
    plt.rcParams['ytick.color'] = '#555555'
    plt.rcParams['font.size'] = 10
    plt.rcParams['axes.labelsize'] = 11
    plt.rcParams['axes.titlesize'] = 12
    plt.rcParams['xtick.labelsize'] = 9
    plt.rcParams['ytick.labelsize'] = 9

apply_premium_style()
# Professional HSL-derived palette (navy, steel blue, teal, sage, grey, light blue, soft orange)
PALETTE = ['#1f4e79', '#4682b4', '#00a896', '#8ecae6', '#778899', '#f4a261', '#e76f51', '#2ec4b6']

# 3. Load Excel Data
df = pd.read_excel(EXCEL_PATH, sheet_name=0)
total_respondents = len(df)

# Data Cleaning and Calculations Helper Functions
def get_freq_pct(series, categories=None, add_total=True, fillna_val=None):
    if fillna_val is not None:
        series = series.fillna(fillna_val)
    counts = series.value_counts()
    
    if categories is None:
        categories = list(counts.index)
    
    data = []
    tot_freq = 0
    for cat in categories:
        freq = counts.get(cat, 0)
        tot_freq += freq
        
    for cat in categories:
        freq = counts.get(cat, 0)
        pct = (freq / tot_freq) * 100 if tot_freq > 0 else 0
        data.append((cat, freq, f"{pct:.1f}%"))
        
    if add_total:
        data.append(("Total", tot_freq, "100.0%"))
        
    return data

# Cleaning town/community names
def clean_location(val):
    if pd.isna(val):
        return "Others"
    val = str(val).strip().lower()
    if "jamestown" in val or "james town" in val:
        return "Jamestown"
    elif "bukom" in val:
        return "Bukom"
    elif "accra" in val:
        return "Accra"
    elif "chorkor" in val:
        return "Chorkor"
    elif "korle gonno" in val:
        return "Korle Gonno"
    else:
        return "Others"

df['cleaned_location'] = df.iloc[:, 10].apply(clean_location) # Col 11 is index 10

# Calculate all tables data
tables_data = {}

# Table 0: Location
t0_cats = ["Jamestown", "Accra", "Bukom", "Chorkor", "Korle Gonno", "Others"]
tables_data[0] = get_freq_pct(df['cleaned_location'], t0_cats)

# Table 1: Gender
t1_cats = ["Male", "Female"]
tables_data[1] = get_freq_pct(df.iloc[:, 2], t1_cats) # Col 3 is index 2

# Table 2: Age Group
ages = df.iloc[:, 3] # Col 4 is index 3
age_groups = []
for age in ages:
    if pd.isna(age):
        age_groups.append(np.nan)
    elif age <= 30:
        age_groups.append("21 - 30") # Including 20 under 21-30 for template match
    elif age <= 40:
        age_groups.append("31 - 40")
    elif age <= 50:
        age_groups.append("41 - 50")
    elif age <= 60:
        age_groups.append("51 - 60")
    else:
        age_groups.append("Above 60")
df['age_group'] = age_groups
t2_cats = ["21 - 30", "31 - 40", "41 - 50", "51 - 60", "Above 60"]
tables_data[2] = get_freq_pct(df['age_group'], t2_cats)

# Table 3: Marital Status
t3_cats = ["Single", "Married", "Divorced", "Separated", "Widowed"]
tables_data[3] = get_freq_pct(df.iloc[:, 4], t3_cats) # Col 5 is index 4

# Table 4: Education level
t4_cats = ["No formal education", "Primary education", "Junior High School", "Senior High School", "Vocational/Technical", "Tertiary education"]
tables_data[4] = get_freq_pct(df.iloc[:, 5], t4_cats) # Col 6 is index 5

# Table 5: Religion
t5_cats = ["Christianity", "Traditional religion", "Islam", "Other (Specify)"]
t5_data = get_freq_pct(df.iloc[:, 6], t5_cats) # Col 7 is index 6
t5_data = [(r[0].replace(" (Specify)", "") if isinstance(r[0], str) else r[0], r[1], r[2]) for r in t5_data]
tables_data[5] = t5_data

# Table 6: Household Size
hh_size = df.iloc[:, 16] # Col 17 is index 16
hh_groups = []
for s in hh_size:
    if pd.isna(s):
        hh_groups.append(np.nan)
    elif s <= 4:
        hh_groups.append("1 - 4")
    elif s <= 7:
        hh_groups.append("5 - 7")
    else:
        hh_groups.append("> 8")
df['hh_group'] = hh_groups
t6_cats = ["1 - 4", "5 - 7", "> 8"]
tables_data[6] = get_freq_pct(df['hh_group'], t6_cats)

# Table 7: Major Occupation
t7_cats = ["Fishing", "Artisanship", "Fish trading", "Farming", "Fish processing", "Other (Specify)"]
t7_data = get_freq_pct(df.iloc[:, 17], t7_cats) # Col 18 is index 17
t7_data = [(r[0].replace(" (Specify)", "") if isinstance(r[0], str) else r[0], r[1], r[2]) for r in t7_data]
tables_data[7] = t7_data

# Table 8: Secondary Occupation
t8_cats = ["Trading", "Casual labour", "Processing", "Farming", "Transport business", "Other (Specify)", "None"]
sec_occ = df.iloc[:, 20].fillna("None") # Col 21 is index 20
t8_data = get_freq_pct(sec_occ, t8_cats)
t8_data = [(r[0].replace(" (Specify)", "") if isinstance(r[0], str) else r[0], r[1], r[2]) for r in t8_data]
tables_data[8] = t8_data

# Table 9: Canoe Ownership (replaced Income Earners)
t9_cats = ["Yes", "No"]
tables_data[9] = get_freq_pct(df.iloc[:, 24], t9_cats) # Col 25 is index 24

# Table 10: Total Household Income
income = df.iloc[:, 67] # Col 68 is index 67
income_groups = []
for inc in income:
    if pd.isna(inc):
        income_groups.append(np.nan)
    elif inc < 3000:
        income_groups.append("< 3000 GHS")
    elif inc <= 3999:
        income_groups.append("3000 - 3999 GHS")
    elif inc <= 4999:
        income_groups.append("4000 - 4999 GHS")
    elif inc <= 5999:
        income_groups.append("5000 - 5999 GHS")
    else:
        income_groups.append("6000 GHS and above")
df['income_group'] = income_groups
t10_cats = ["< 3000 GHS", "3000 - 3999 GHS", "4000 - 4999 GHS", "5000 - 5999 GHS", "6000 GHS and above"]
tables_data[10] = get_freq_pct(df['income_group'], t10_cats)

# Table 11: Employment Status in Fishing (replaced No of Dependents)
t11_cats = ["Both owner and crew member", "Crew member", "Canoe owner"]
tables_data[11] = get_freq_pct(df.iloc[:, 23], t11_cats) # Col 24 is index 23

# Table 12: Associations/FBO
t12_cats = ["Yes", "No"]
tables_data[12] = get_freq_pct(df.iloc[:, 27], t12_cats) # Col 28 is index 27

# Table 13: Savings Status (replaced Applied for Credit)
t13_cats = ["Yes", "No"]
tables_data[13] = get_freq_pct(df.iloc[:, 76], t13_cats) # Col 77 is index 76

# Table 14: Reinvestment Status (replaced Received Credit)
t14_cats = ["Yes", "No"]
tables_data[14] = get_freq_pct(df.iloc[:, 77], t14_cats) # Col 78 is index 77

# Table 15: Type of Association (replaced Source of Credit)
assoc_type = df.iloc[:, 28].fillna("None") # Col 29 is index 28
t15_cats = ["Fisher association", "Cooperative society", "Fish processors association", "Savings group", "Religious association", "Other (Specify)", "None"]
t15_data = get_freq_pct(assoc_type, t15_cats)
t15_data = [(r[0].replace(" (Specify)", "") if isinstance(r[0], str) else r[0], r[1], r[2]) for r in t15_data]
tables_data[15] = t15_data

# Table 16: Average Weekly Expenditure on Fuel (replaced Credit Amount GHS)
fuel_exp = df.iloc[:, 33] # Col 34 is index 33
fuel_groups = []
for f in fuel_exp:
    if pd.isna(f):
        fuel_groups.append("Not Applicable")
    elif f < 500:
        fuel_groups.append("< 500 GHS")
    elif f <= 699:
        fuel_groups.append("500 - 699 GHS")
    elif f <= 899:
        fuel_groups.append("700 - 899 GHS")
    else:
        fuel_groups.append("900 GHS and above")
df['fuel_group'] = fuel_groups
t16_cats = ["< 500 GHS", "500 - 699 GHS", "700 - 899 GHS", "900 GHS and above", "Not Applicable"]
tables_data[16] = get_freq_pct(df['fuel_group'], t16_cats)

# Table 17: Descriptive Statistics of Fuel Expenditure (replaced Mean/Median/Min/Max of Credit Amount)
fuel_valid = fuel_exp.dropna()
tables_data[17] = [
    (
        f"{fuel_valid.mean():,.2f}",
        f"{fuel_valid.median():,.2f}",
        f"{fuel_valid.min():,.2f}",
        f"{fuel_valid.max():,.2f}",
        f"{fuel_valid.std():,.2f}"
    )
]

# Table 18: Rating of Fuel Price Increase (replaced Interest Rate %)
t18_cats = ["Very high", "High", "Moderate", "Low", "Very low", "N/A"]
fuel_inc = df.iloc[:, 34].fillna("N/A") # Col 35 is index 34
tables_data[18] = get_freq_pct(fuel_inc, t18_cats)

# Table 19: Frequency of Fuel Shortages (replaced Repayment Period)
t19_cats = ["Very often", "Often", "Occasionally", "Rarely", "N/A"]
fuel_sh = df.iloc[:, 36].fillna("N/A") # Col 37 is index 36
tables_data[19] = get_freq_pct(fuel_sh, t19_cats)

# Table 20: Years of Fishing Experience (replaced Poultry Experience)
exp = df.iloc[:, 22] # Col 23 is index 22
exp_groups = []
for e in exp:
    if pd.isna(e):
        exp_groups.append(np.nan)
    elif e < 10:
        exp_groups.append("< 10 years")
    elif e < 20:
        exp_groups.append("10 - 19 years")
    else:
        exp_groups.append("20 - 30 years")
df['exp_group'] = exp_groups
t20_cats = ["< 10 years", "10 - 19 years", "20 - 30 years"]
tables_data[20] = get_freq_pct(df['exp_group'], t20_cats)

# Table 21: Weekly Fish Catch Bracket (replaced Flock Size)
catch = df.iloc[:, 65] # Col 66 is index 65
catch_groups = []
for c in catch:
    if pd.isna(c):
        catch_groups.append(np.nan)
    elif c < 150:
        catch_groups.append("< 150 kg")
    elif c <= 249:
        catch_groups.append("150 - 249 kg")
    elif c <= 349:
        catch_groups.append("250 - 349 kg")
    else:
        catch_groups.append("350 kg and above")
df['catch_group'] = catch_groups
t21_cats = ["< 150 kg", "150 - 249 kg", "250 - 349 kg", "350 kg and above"]
tables_data[21] = get_freq_pct(df['catch_group'], t21_cats)

# Table 22: Belief that Cooperative Groups Can Improve Welfare (replaced Poultry Training)
t22_cats = ["Yes", "No"]
tables_data[22] = get_freq_pct(df.iloc[:, 93], t22_cats) # Col 94 is index 93

# Table 23: Institutional Factors (calculated from 8 columns)
inst_cols = [
    ("Belong to community association", df.iloc[:, 27]), # Col 28
    ("Own canoe", df.iloc[:, 24]), # Col 25
    ("Experience high fuel costs", df.iloc[:, 32]), # Col 33
    ("Experience premix fuel shortages", df.iloc[:, 35]), # Col 36
    ("Experience declining catch", df.iloc[:, 37]), # Col 38
    ("Experience frequent gear damage", df.iloc[:, 42]), # Col 43
    ("Experience coastal pollution", df.iloc[:, 46]), # Col 47
    ("Keep income and expense records", df.iloc[:, 74]) # Col 75
]
t23_rows = []
for label, col in inst_cols:
    val_counts = col.value_counts()
    yes_c = val_counts.get("Yes", 0)
    no_c = val_counts.get("No", 0)
    tot = yes_c + no_c
    yes_pct = (yes_c / tot * 100) if tot > 0 else 0
    no_pct = (no_c / tot * 100) if tot > 0 else 0
    t23_rows.append((
        label,
        f"{yes_c} ({yes_pct:.1f}%)",
        f"{no_c} ({no_pct:.1f}%)",
        f"{tot} (100.0%)"
    ))
tables_data[23] = t23_rows

# Table 24: Distribution of Respondents Based on Selling Price (replaced Revenue Growth Bracket)
price = df.iloc[:, 66] # Col 67 is index 66
price_groups = []
for p in price:
    if pd.isna(p):
        price_groups.append(np.nan)
    elif p < 30:
        price_groups.append("Extremely Low (< 30 GHS)")
    elif p <= 44:
        price_groups.append("Very Low (30 to 44 GHS)")
    elif p <= 59:
        price_groups.append("Low (45 to 59 GHS)")
    elif p <= 74:
        price_groups.append("Moderate (60 to 74 GHS)")
    elif p <= 89:
        price_groups.append("High (75 to 89 GHS)")
    elif p <= 104:
        price_groups.append("Very High (90 to 104 GHS)")
    else:
        price_groups.append("Extremely High (105 GHS and above)")
df['price_group'] = price_groups
t24_cats = [
    "Extremely Low (< 30 GHS)",
    "Very Low (30 to 44 GHS)",
    "Low (45 to 59 GHS)",
    "Moderate (60 to 74 GHS)",
    "High (75 to 89 GHS)",
    "Very High (90 to 104 GHS)",
    "Extremely High (105 GHS and above)"
]
tables_data[24] = get_freq_pct(df['price_group'], t24_cats)

# Table 25: Income Sufficiency Category (replaced Scalability Category)
rating = df.iloc[:, 81] # Col 82 is index 81
rating_groups = []
for r in rating:
    if pd.isna(r):
        rating_groups.append(np.nan)
    elif r in ["Good", "Very good"]:
        rating_groups.append("Good/Very Good (Sufficient)")
    elif r == "Average":
        rating_groups.append("Average (Stagnant)")
    else:
        rating_groups.append("Poor/Very Poor (Declining)")
df['rating_group'] = rating_groups
t25_cats = ["Good/Very Good (Sufficient)", "Average (Stagnant)", "Poor/Very Poor (Declining)"]
tables_data[25] = get_freq_pct(df['rating_group'], t25_cats)

# Table 26: Descriptive Statistics of Fish Selling Price (replaced Descriptive Stats of Revenue Growth)
price_valid = price.dropna()
tables_data[26] = [
    (
        f"{price_valid.mean():,.2f}",
        f"{price_valid.median():,.2f}",
        f"{price_valid.min():,.2f}",
        f"{price_valid.max():,.2f}",
        f"{price_valid.std():,.2f}"
    )
]

# Table 27: Gear Repair Expenditure Bracket (replaced Initial Capital Invested)
gear = df.iloc[:, 45] # Col 46 is index 45
gear_groups = []
for g in gear:
    if pd.isna(g):
        gear_groups.append("Not Applicable")
    elif g < 300:
        gear_groups.append("< 300 GHS")
    elif g <= 399:
        gear_groups.append("300 - 399 GHS")
    elif g <= 499:
        gear_groups.append("400 - 499 GHS")
    else:
        gear_groups.append("500 GHS and above")
df['gear_group'] = gear_groups
t27_cats = ["< 300 GHS", "300 - 399 GHS", "400 - 499 GHS", "500 GHS and above", "Not Applicable"]
tables_data[27] = get_freq_pct(df['gear_group'], t27_cats)

# Table 28: Weekly Fish Catch Bracket (repeat catch)
tables_data[28] = tables_data[21]

# Table 29: Fuel Expenditure Bracket (repeat fuel but exclude N/A)
fuel_active = fuel_exp.dropna()
df_active_fuel = pd.DataFrame({'fuel': fuel_active})
active_groups = []
for f in df_active_fuel['fuel']:
    if f < 500:
        active_groups.append("< 500 GHS")
    elif f <= 699:
        active_groups.append("500 - 699 GHS")
    elif f <= 899:
        active_groups.append("700 - 899 GHS")
    else:
        active_groups.append("900 GHS and above")
df_active_fuel['group'] = active_groups
t29_cats = ["< 500 GHS", "500 - 699 GHS", "700 - 899 GHS", "900 GHS and above"]
tables_data[29] = get_freq_pct(df_active_fuel['group'], t29_cats)

# Table 30: Household Size Group (repeat hh_size but with 4 categories)
hh_4cats = []
for s in hh_size:
    if pd.isna(s):
        hh_4cats.append(np.nan)
    elif s <= 2:
        hh_4cats.append("1 - 2")
    elif s <= 4:
        hh_4cats.append("3 - 4")
    elif s <= 6:
        hh_4cats.append("5 - 6")
    else:
        hh_4cats.append("7 and above")
df['hh_4cat'] = hh_4cats
t30_cats = ["1 - 2", "3 - 4", "5 - 6", "7 and above"]
tables_data[30] = get_freq_pct(df['hh_4cat'], t30_cats)

# Table 31: Years of Fishing Experience Bracket (alternative split)
exp_5cats = []
for e in exp:
    if pd.isna(e):
        exp_5cats.append(np.nan)
    elif e < 10:
        exp_5cats.append("< 10 years")
    elif e <= 14:
        exp_5cats.append("10 - 14 years")
    elif e <= 19:
        exp_5cats.append("15 - 19 years")
    elif e <= 24:
        exp_5cats.append("20 - 24 years")
    else:
        exp_5cats.append("25 - 30 years")
df['exp_5cat'] = exp_5cats
t31_cats = ["< 10 years", "10 - 14 years", "15 - 19 years", "20 - 24 years", "25 - 30 years"]
tables_data[31] = get_freq_pct(df['exp_5cat'], t31_cats)

# Table 32: Access to Cold Storage Facilities (replaced Input Access)
t32_cats = ["Yes", "No"]
tables_data[32] = get_freq_pct(df.iloc[:, 48], t32_cats) # Col 49 is index 48

# Table 33: Reinvests Revenue into Fishing Operations
t33_cats = ["Yes", "No"]
tables_data[33] = get_freq_pct(df.iloc[:, 77], t33_cats) # Col 78 is index 77

# Table 34: Overall Rating of Current Fishing Income
t34_cats = ["Very poor", "Poor", "Average", "Good", "Very good"]
tables_data[34] = get_freq_pct(df.iloc[:, 81], t34_cats) # Col 82 is index 81

# Table 35: Rating of Cost of Fishing Equipment and Inputs
t35_cats = ["Very low", "Low", "Moderate", "High", "Very high"]
tables_data[35] = get_freq_pct(df.iloc[:, 54], t35_cats) # Col 55 is index 54

# Garrett's Ranking of Operational Challenges
rank_cols = [c for c in df.columns if 'Rank for:' in c]
garrett_scores_map = {1: 77, 2: 63, 3: 54, 4: 46, 5: 37, 6: 23}
df_garrett_scores = df[rank_cols].map(lambda x: garrett_scores_map.get(x, 0) if pd.notna(x) else 0)
mean_garrett_scores = df_garrett_scores.mean()
sorted_garrett = mean_garrett_scores.sort_values(ascending=False)
challenges_clean = {
    "Rank for: Fuel cost": "Fuel cost",
    "Rank for: Gear damage": "Gear damage",
    "Rank for: Pollution": "Pollution",
    "Rank for: Credit constraints": "Credit constraints",
    "Rank for: Fish stock decline": "Fish stock decline",
    "Rank for: Inadequate storage facilities": "Inadequate storage facilities"
}
garrett_data = []
for rank_idx, (col_name, score) in enumerate(sorted_garrett.items()):
    clean_name = challenges_clean.get(col_name, col_name)
    rank_str = ["Rank I", "Rank II", "Rank III", "Rank IV", "Rank V", "Rank VI"][rank_idx]
    garrett_data.append((clean_name, f"{score:.2f}", rank_str))

print("All tables data calculated successfully!")

# 4. Generate the 20 Custom Charts
def save_chart(filename, title, data_list, chart_type='bar', color_idx=0):
    fig, ax = plt.subplots(figsize=(6.5, 3.8), dpi=200)
    
    plot_data = [r for r in data_list if r[0] != "Total"]
    labels = [str(r[0]) for r in plot_data]
    
    freqs = []
    for r in plot_data:
        if len(r) >= 2:
            freqs.append(float(r[1]))
        else:
            freqs.append(0)
            
    pcts = []
    for r in plot_data:
        if len(r) >= 3:
            pct_val = r[2].replace('%', '')
            pcts.append(float(pct_val))
        else:
            pcts.append(0)
            
    color = PALETTE[color_idx % len(PALETTE)]
    
    if chart_type == 'bar':
        bars = ax.bar(labels, freqs, color=color, width=0.55, edgecolor='#2c3e50', linewidth=0.5)
        ax.set_ylabel('Frequency')
        ax.grid(axis='y', linestyle='--', alpha=0.5)
        for bar, pct in zip(bars, pcts):
            height = bar.get_height()
            ax.annotate(f'{int(height)}\n({pct:.1f}%)',
                        xy=(bar.get_x() + bar.get_width() / 2, height),
                        xytext=(0, 3),
                        textcoords="offset points",
                        ha='center', va='bottom', fontsize=8, color='#2c3e50')
        ymax = max(freqs) * 1.18 if freqs else 10
        ax.set_ylim(0, ymax)
        
    elif chart_type == 'barh':
        y_pos = np.arange(len(labels))
        bars = ax.barh(y_pos, freqs, color=color, height=0.55, edgecolor='#2c3e50', linewidth=0.5)
        ax.set_yticks(y_pos)
        ax.set_yticklabels(labels)
        ax.invert_yaxis()
        ax.set_xlabel('Frequency')
        ax.grid(axis='x', linestyle='--', alpha=0.5)
        for bar, pct in zip(bars, pcts):
            width = bar.get_width()
            ax.annotate(f' {int(width)} ({pct:.1f}%)',
                        xy=(width, bar.get_y() + bar.get_height() / 2),
                        xytext=(3, 0),
                        textcoords="offset points",
                        ha='left', va='center', fontsize=8, color='#2c3e50')
        xmax = max(freqs) * 1.18 if freqs else 10
        ax.set_xlim(0, xmax)
        
    elif chart_type == 'donut':
        colors = [PALETTE[i % len(PALETTE)] for i in range(len(labels))]
        wedges, texts, autotexts = ax.pie(freqs, labels=labels, autopct='%1.1f%%',
                                          startangle=90, colors=colors, 
                                          textprops=dict(color="#2c3e50", fontsize=9),
                                          wedgeprops=dict(width=0.4, edgecolor='#ffffff', linewidth=1.5))
        for autotext in autotexts:
            autotext.set_fontsize(9)
            autotext.set_weight('bold')
            autotext.set_color('white')
            
    ax.set_title(title, pad=15, weight='bold', color='#1f4e79')
    plt.tight_layout()
    chart_path = os.path.join(CHARTS_DIR, filename)
    plt.savefig(chart_path, dpi=200)
    plt.close()
    return chart_path

# Generate the 20 charts
charts_paths = {}

charts_paths[0] = save_chart("chart_0_location.png", "Distribution of Respondents by Location", tables_data[0], 'barh', 0)
charts_paths[1] = save_chart("chart_1_gender.png", "Gender Distribution of Respondents", tables_data[1], 'bar', 1)
charts_paths[2] = save_chart("chart_2_age.png", "Age Distribution of Respondents", tables_data[2], 'bar', 2)
charts_paths[3] = save_chart("chart_3_marital.png", "Marital Status of Respondents", tables_data[3], 'bar', 3)
charts_paths[4] = save_chart("chart_4_education.png", "Educational Attainment of Respondents", tables_data[4], 'barh', 4)
charts_paths[5] = save_chart("chart_5_religion.png", "Religious Beliefs of Respondents", tables_data[5], 'bar', 5)
charts_paths[6] = save_chart("chart_6_hh_size.png", "Household Size of Respondents", tables_data[6], 'bar', 6)
charts_paths[7] = save_chart("chart_7_major_occ.png", "Major Occupation of Respondents", tables_data[7], 'barh', 7)
charts_paths[8] = save_chart("chart_8_secondary_occ.png", "Secondary Occupation of Respondents", tables_data[8], 'barh', 0)
charts_paths[9] = save_chart("chart_9_canoe_ownership.png", "Canoe Ownership Status", tables_data[9], 'bar', 1)
charts_paths[10] = save_chart("chart_10_hh_income.png", "Average Monthly Household Income", tables_data[10], 'bar', 2)
charts_paths[11] = save_chart("chart_11_employment_status.png", "Employment Status in Fishing", tables_data[11], 'barh', 3)
charts_paths[12] = save_chart("chart_12_association.png", "Community-Based Association Membership", tables_data[12], 'bar', 4)
charts_paths[13] = save_chart("chart_13_savings.png", "Savings Status of Respondents", tables_data[13], 'bar', 5)
charts_paths[14] = save_chart("chart_14_reinvestment.png", "Reinvestment of Fishing Revenue", tables_data[14], 'bar', 6)
charts_paths[15] = save_chart("chart_15_assoc_type.png", "Type of Association Belonged To", tables_data[15], 'barh', 7)
charts_paths[16] = save_chart("chart_16_fuel_exp.png", "Average Weekly Expenditure on Fuel Bracket", tables_data[16], 'bar', 0)
charts_paths[17] = save_chart("chart_17_fuel_price_rating.png", "Rating of Fuel Price Increase Over the Years", tables_data[18], 'bar', 1)
charts_paths[18] = save_chart("chart_18_fuel_shortages_freq.png", "Frequency of Fuel Shortages", tables_data[19], 'bar', 2)
charts_paths[19] = save_chart("chart_19_selling_price.png", "Average Selling Price of Fish per kg", tables_data[24], 'barh', 3)

# Drawing 20: Garrett's Ranking Chart
fig, ax = plt.subplots(figsize=(6.5, 3.8), dpi=200)
challenges_chart = [r[0] for r in garrett_data]
scores_chart = [float(r[1]) for r in garrett_data]
colors = [PALETTE[i % len(PALETTE)] for i in range(len(challenges_chart))]
bars = ax.bar(challenges_chart, scores_chart, color=colors, width=0.55, edgecolor='#2c3e50', linewidth=0.5)
ax.set_ylabel('Mean Garrett Score')
ax.grid(axis='y', linestyle='--', alpha=0.5)
for bar in bars:
    height = bar.get_height()
    ax.annotate(f'{height:.2f}',
                xy=(bar.get_x() + bar.get_width() / 2, height),
                xytext=(0, 3),
                textcoords="offset points",
                ha='center', va='bottom', fontsize=8, color='#2c3e50')
ax.set_ylim(0, 100)
ax.set_title("Garrett's Ranking of Operational Challenges", pad=15, weight='bold', color='#1f4e79')
plt.xticks(rotation=15, ha='right')
plt.tight_layout()
chart_path_garrett = os.path.join(CHARTS_DIR, "chart_20_garrett_ranking.png")
plt.savefig(chart_path_garrett, dpi=200)
plt.close()
charts_paths[20] = chart_path_garrett

print("All 21 premium charts generated and saved successfully!")

# 5. Backup the Word Document
shutil.copy2(DOCX_PATH, r"C:\Users\Hughe\Downloads\survey_auto_filler\Data Analysis_Backup.docx")
print("Backed up original Word document to Data Analysis_Backup.docx")

# 6. Read and Update Word Document
doc = docx.Document(DOCX_PATH)

# Function to update table rows
def update_word_table(table, data_rows):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
        
    for row_data in data_rows:
        new_row = table.add_row()
        for col_idx, val in enumerate(row_data):
            if col_idx < len(new_row.cells):
                cell = new_row.cells[col_idx]
                cell.text = str(val)
                p = cell.paragraphs[0]
                if col_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Apply table updates
for t_idx, data in tables_data.items():
    if t_idx < len(doc.tables):
        update_word_table(doc.tables[t_idx], data)
        print(f"Updated Table {t_idx} (Header: {[c.text for c in doc.tables[t_idx].rows[0].cells[:2]]})")

# Define Headings Renaming Mapping
heading_renames = {
    "Other Occupation": "Secondary Occupation of Respondents",
    "Income Earners in the Household of Respondents": "Canoe Ownership of Respondents",
    "Total Household Income of Respondents": "Average Monthly Household Income of Respondents",
    "Dependents in Households of Respondents": "Employment Status in Fishing of Respondents",
    "Nature of Credit": "Financial Characteristics of Respondents",
    "Applied for Credit": "Savings Status of Respondents",
    "Received Credit": "Reinvestment of Fishing Revenue",
    "Main Source of Credit": "Type of Community Association Belonged To",
    "Amount of Credit Received": "Average Weekly Expenditure on Fuel",
    "Interest Rate": "Rating of Fuel Price Increase Over the Years",
    "Repayment Period": "Frequency of Fuel Shortages Affecting Fishing",
    "Factors Influencing Access to Credit": "Key Operational Challenges and Support",
    "Technical Factors": "Fishing Experience and Capacity",
    "Institutional Factors": "Experience of Operational Challenges and Records",
    "Nature of Farm Scalability": "Financial Performance and Input Cost Analysis",
    "Distribution of Farms Based on their Revenue Growth Rate": "Distribution of Respondents Based on Average Selling Price of Fish per kg",
    "Influence of Access to Credit on Farm Scalability": "Influence of Operational Costs on Fishing Performance",
    "Initial Capital Invested (GHS)": "Average Monthly Expenditure on Gear Repair/Maintenance (GHS)",
    "Current Flock Size": "Average Weekly Fish Catch Bracket (kg)",
    "Average Feed Cost per Production Cycle(GHS)": "Average Weekly Expenditure on Fuel Bracket (GHS)",
    "Number of Workers Employed": "Household Size of Respondents (Alternative Grouping)",
    "Number of Production Cycles per Year": "Years of Fishing Experience Bracket",
    "Access to Poultry Inputs": "Access to Cold Storage Facilities",
    "Reinvestment of Farm Revenue": "Reinvestment of Fishing Revenue",
    "Percentage of Revenue Reinvested": "Overall Rating of Current Fishing Income",
    "Quality of Feed Used": "Rating of the Cost of Fishing Equipment and Inputs"
}

# Apply paragraph renames
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in heading_renames:
        p.text = heading_renames[txt]
        for run in p.runs:
            run.font.name = 'Arial'
            run.bold = True
        print(f"Renamed heading: '{txt}' -> '{heading_renames[txt]}'")

# Apply Table Headers Renaming
table_headers = {
    7: ["Major Occupation", "Frequency", "Percentage %"],
    8: ["Secondary Occupation", "Frequency", "Percentage %"],
    9: ["Canoe Ownership", "Frequency", "Percentage %"],
    10: ["Income Bracket (GHS)", "Frequency", "Percentage %"],
    11: ["Employment Status", "Frequency", "Percentage %"],
    13: ["Saves Part of Income", "Frequency", "Percentage %"],
    14: ["Reinvests in Fishing", "Frequency", "Percentage %"],
    15: ["Association Type", "Frequency", "Percentage %"],
    16: ["Weekly Fuel Cost (GHS)", "Frequency", "Percentage %"],
    18: ["Fuel Price Increase Rating", "Frequency", "Percentage %"],
    19: ["Fuel Shortage Frequency", "Frequency", "Percentage %"],
    20: ["Fishing Experience (Years)", "Frequency", "Percentage %"],
    21: ["Weekly Catch Bracket (kg)", "Frequency", "Percentage %"],
    22: ["Believes Cooperatives Help", "Frequency", "Percentage %"],
    24: ["Range", "Selling Price Bracket (GHS)", "Frequency", "Percentage %"],
    25: ["Income Sufficiency Category", "Frequency", "Percentage %"],
    26: ["Mean Price (GHS)", "Median Price (GHS)", "Minimum Price (GHS)", "Maximum Price (GHS)", "Standard Deviation (GHS)"],
    27: ["Gear Repair Expenditure (GHS)", "Frequency", "Percentage %"],
    28: ["Weekly Catch Bracket (kg)", "Frequency", "Percentage %"],
    29: ["Weekly Fuel Cost (GHS)", "Frequency", "Percentage %"],
    30: ["Household Size Group", "Frequency", "Percentage %"],
    31: ["Fishing Experience (Years)", "Frequency", "Percentage %"],
    32: ["Cold Storage Access", "Frequency", "Percentage %"],
    33: ["Reinvests Fishing Revenue", "Frequency", "Percentage %"],
    34: ["Current Income Rating", "Frequency", "Percentage %"],
    35: ["Equipment Cost Rating", "Frequency", "Percentage %"]
}

for t_idx, headers in table_headers.items():
    if t_idx < len(doc.tables):
        table = doc.tables[t_idx]
        for col_idx, h_text in enumerate(headers):
            if col_idx < len(table.rows[0].cells):
                table.rows[0].cells[col_idx].text = h_text
                for p in table.rows[0].cells[col_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
        print(f"Renamed Table {t_idx} headers to {headers}")

# Replacing Drawings (Charts)
drawing_idx = 0
for idx, element in enumerate(doc.element.body):
    tag = element.tag.split('}')[-1]
    if tag == 'p':
        p = docx.text.paragraph.Paragraph(element, doc)
        drawings_in_p = element.xpath('.//w:drawing')
        if drawings_in_p:
            if drawing_idx in charts_paths:
                p.text = ""  # Clear drawing and run elements
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                chart_img = charts_paths[drawing_idx]
                run.add_picture(chart_img, width=Inches(5.5))
                print(f"Replaced Drawing {drawing_idx} (Para element {idx}) with {chart_img}")
            drawing_idx += 1

# Insert Garrett's Ranking Section right after Table 23
tbl23 = doc.tables[23]._tbl
p_garrett = doc.add_paragraph()
p_garrett.text = "Garrett's Ranking of Operational Challenges"
p_garrett.style = 'Heading 2'
for run in p_garrett.runs:
    run.font.name = 'Arial'
    run.bold = True
tbl23.addnext(p_garrett._p)

# Insert Table right after Heading paragraph
table_garrett = doc.add_table(rows=7, cols=3)
table_garrett.style = doc.tables[23].style
p_garrett._p.addnext(table_garrett._tbl)

# Populate Garrett Table
headers = ['Operational Challenge', 'Mean Garrett Score', 'Final Rank']
for col_idx, h in enumerate(headers):
    table_garrett.rows[0].cells[col_idx].text = h
    for p_cell in table_garrett.rows[0].cells[col_idx].paragraphs:
        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_cell.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            
for r_idx, (challenge, score, rank) in enumerate(garrett_data):
    row = table_garrett.rows[r_idx + 1]
    row.cells[0].text = challenge
    row.cells[1].text = score
    row.cells[2].text = rank
    
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for cell in row.cells:
        for p_cell in cell.paragraphs:
            for run in p_cell.runs:
                run.font.name = 'Arial'
                
# Insert Chart right after Table
p_chart = doc.add_paragraph()
p_chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_chart = p_chart.add_run()
run_chart.add_picture(charts_paths[20], width=Inches(5.5))
table_garrett._tbl.addnext(p_chart._p)

# Save the final file
doc.save(DOCX_OUT_PATH)
try:
    shutil.copy2(DOCX_OUT_PATH, DOCX_PATH)
    print(f"Also copied/overwrote original template at: {DOCX_PATH}")
except PermissionError:
    print(f"\n[WARNING] Could not overwrite '{DOCX_PATH}' because it is open in another program (e.g. MS Word). Please close it and copy '{DOCX_OUT_PATH}' manually.")

print(f"\nSUCCESS! Document successfully created at: {DOCX_OUT_PATH}")
